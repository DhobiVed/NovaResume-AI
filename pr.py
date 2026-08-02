# import streamlit as st
# from groq import Groq
# import time
# from datetime import datetime

# # ===== API KEY =====
# GROQ_API_KEY = "your_groq_api_key_here"
# client = Groq(api_key=GROQ_API_KEY)

# # ===== CONFIG =====
# st.set_page_config(
#     page_title="NovaChat",
#     page_icon="💬",
#     layout="wide"
# )

# # ===== SESSION STATE =====
# if "chats" not in st.session_state:
#     cid = str(int(time.time()))
#     st.session_state.chats = {
#         cid: {
#             "title": "New Chat",
#             "created": datetime.now(),
#             "messages": [
#                 {"role": "system", "content": "You are a helpful AI assistant."}
#             ],
#         }
#     }
#     st.session_state.current = cid

# # ===== FUNCTIONS =====
# def new_chat():
#     cid = str(int(time.time()))
#     st.session_state.chats[cid] = {
#         "title": "New Chat",
#         "created": datetime.now(),
#         "messages": [
#             {"role": "system", "content": "You are a helpful AI assistant."}
#         ],
#     }
#     st.session_state.current = cid

# def add_msg(role, content):
#     chat = st.session_state.chats[st.session_state.current]
#     chat["messages"].append({"role": role, "content": content})

#     if role == "user" and len(chat["messages"]) == 2:
#         chat["title"] = content[:30]

# # ===== CSS (Clean ChatGPT Style) =====
# st.markdown("""
# <style>
# .stApp {background:#0f172a; color:white;}

# .chat-header {
#     text-align:center;
#     font-size:28px;
#     font-weight:700;
#     margin-bottom:10px;
# }

# .chat-message {
#     display:flex;
#     margin:12px 0;
#     animation:fadeIn .25s ease-in;
# }

# @keyframes fadeIn {
#     from {opacity:0; transform:translateY(5px);}
#     to {opacity:1; transform:translateY(0);}
# }

# .user {justify-content:flex-end;}
# .assistant {justify-content:flex-start;}

# .bubble {
#     max-width:70%;
#     padding:14px 18px;
#     border-radius:18px;
#     font-size:15px;
#     line-height:1.5;
# }

# .user-bubble {
#     background:#2563eb;
#     border-bottom-right-radius:6px;
# }

# .assistant-bubble {
#     background:#1e293b;
#     color:#e5e7eb;
#     border-bottom-left-radius:6px;
# }

# section[data-testid="stSidebar"] {
#     background:#020617;
# }

# .stChatInput input {
#     border-radius:20px !important;
#     padding:14px !important;
# }
# </style>
# """, unsafe_allow_html=True)

# # ===== SIDEBAR =====
# with st.sidebar:
#     st.title("💬 NovaChat")

#     if st.button("➕ New Chat", use_container_width=True):
#         new_chat()
#         st.rerun()

#     st.divider()

#     for cid, data in sorted(
#         st.session_state.chats.items(),
#         key=lambda x: x[1]["created"],
#         reverse=True
#     ):
#         if st.button(data["title"], key=cid, use_container_width=True):
#             st.session_state.current = cid
#             st.rerun()

#     st.divider()

#     model = st.selectbox(
#         "Model",
#         ["llama-3.1-8b-instant",
#          "llama-3.3-70b-versatile"]
#     )

# # ===== MAIN =====
# chat = st.session_state.chats[st.session_state.current]

# st.markdown(
#     f"<div class='chat-header'>✨ {chat['title']}</div>",
#     unsafe_allow_html=True
# )

# # ===== SHOW MESSAGES =====
# for m in chat["messages"][1:]:
#     if m["role"] == "user":
#         st.markdown(f"""
#         <div class="chat-message user">
#             <div class="bubble user-bubble">{m["content"]}</div>
#         </div>
#         """, unsafe_allow_html=True)
#     else:
#         st.markdown(f"""
#         <div class="chat-message assistant">
#             <div class="bubble assistant-bubble">{m["content"]}</div>
#         </div>
#         """, unsafe_allow_html=True)

# # ===== INPUT =====
# prompt = st.chat_input("Message Nova...")

# if prompt:
#     add_msg("user", prompt)

#     st.markdown(f"""
#     <div class="chat-message user">
#         <div class="bubble user-bubble">{prompt}</div>
#     </div>
#     """, unsafe_allow_html=True)

#     placeholder = st.empty()
#     reply = ""

#     stream = client.chat.completions.create(
#         model=model,
#         messages=chat["messages"],
#         stream=True,
#     )

#     for chunk in stream:
#         if chunk.choices[0].delta.content:
#             reply += chunk.choices[0].delta.content
#             placeholder.markdown(f"""
#             <div class="chat-message assistant">
#                 <div class="bubble assistant-bubble">{reply}</div>
#             </div>
#             """, unsafe_allow_html=True)

#     add_msg("assistant", reply)

# st.markdown(
#     "<center style='opacity:.5;margin-top:20px'>Made with ❤️ using Groq + Streamlit</center>",
#     unsafe_allow_html=True
# )








# import streamlit as st
# from groq import Groq
# import time
# from datetime import datetime
# import json
# import os
# import tempfile
# import pdfplumber
# from docx import Document
# from fpdf import FPDF
# import base64

# # ===== API KEY =====
# GROQ_API_KEY = "your_groq_api_key_here"
# client = Groq(api_key=GROQ_API_KEY)

# # ===== CONFIG =====
# st.set_page_config(
#     page_title="Nova Advanced",
#     page_icon="✨",
#     layout="wide"
# )

# # ===== SESSION STATE INIT =====
# if "chats" not in st.session_state:
#     cid = str(int(time.time()))
#     st.session_state.chats = {
#         cid: {
#             "title": "New Chat",
#             "created": datetime.now().isoformat(),
#             "system_prompt": "You are Nova, a highly accurate and factual AI assistant. Think step-by-step. Give concise, truthful answers. If unsure, say you don't know.",
#             "messages": [],  # each message: {"role": str, "content": str}
#             "memory": {"user_name": None}  # simple memory
#         }
#     }
#     st.session_state.current = cid
#     st.session_state.stop_generation = False
#     st.session_state.regenerate_target = None
#     st.session_state.theme = "dark"
#     st.session_state.user_name = "User"

# # ===== HELPER FUNCTIONS =====
# def new_chat():
#     cid = str(int(time.time()))
#     st.session_state.chats[cid] = {
#         "title": "New Chat",
#         "created": datetime.now().isoformat(),
#         "system_prompt": "You are Nova, a highly accurate and factual AI assistant. Think step-by-step. Give concise, truthful answers. If unsure, say you don't know.",
#         "messages": [],
#         "memory": {"user_name": st.session_state.user_name}
#     }
#     st.session_state.current = cid

# def delete_chat(cid):
#     if cid in st.session_state.chats:
#         del st.session_state.chats[cid]
#     if st.session_state.current == cid:
#         if st.session_state.chats:
#             st.session_state.current = next(iter(st.session_state.chats))
#         else:
#             new_chat()
#     st.rerun()

# def rename_chat(cid, new_title):
#     if cid in st.session_state.chats:
#         st.session_state.chats[cid]["title"] = new_title

# def add_msg(role, content):
#     chat = st.session_state.chats[st.session_state.current]
#     chat["messages"].append({"role": role, "content": content})
#     if role == "user" and len(chat["messages"]) == 1:
#         chat["title"] = content[:30]

# def update_system_prompt(new_prompt):
#     st.session_state.chats[st.session_state.current]["system_prompt"] = new_prompt

# def get_messages_for_api():
#     """Return messages list without any extra fields (only role, content)"""
#     chat = st.session_state.chats[st.session_state.current]
#     msgs = [{"role": "system", "content": chat["system_prompt"]}]
#     for m in chat["messages"]:
#         msgs.append({"role": m["role"], "content": m["content"]})
#     return msgs

# def export_chat_as_markdown(cid):
#     chat = st.session_state.chats[cid]
#     lines = [f"# {chat['title']}\n", f"*Created: {chat['created']}*\n"]
#     for msg in chat["messages"]:
#         role = "**User**" if msg["role"] == "user" else "**Nova**"
#         lines.append(f"\n{role}: {msg['content']}\n")
#     return "\n".join(lines)

# def export_chat_as_pdf(cid):
#     chat = st.session_state.chats[cid]
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt=chat['title'], ln=1, align='C')
#     pdf.ln(10)
#     for msg in chat["messages"]:
#         role = "User: " if msg["role"] == "user" else "Nova: "
#         pdf.multi_cell(0, 10, txt=role + msg['content'])
#         pdf.ln(5)
#     return pdf.output(dest='S').encode('latin1')

# def extract_text_from_file(uploaded_file):
#     text = ""
#     if uploaded_file.type == "application/pdf":
#         with pdfplumber.open(uploaded_file) as pdf:
#             for page in pdf.pages:
#                 text += page.extract_text() or ""
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         doc = Document(uploaded_file)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     elif uploaded_file.type == "text/plain":
#         text = uploaded_file.read().decode("utf-8")
#     return text

# # ===== THEME CSS =====
# def get_theme_css():
#     if st.session_state.theme == "dark":
#         return """
#         <style>
#         .stApp { background: linear-gradient(135deg, #0f172a, #020617); color: white; }
#         .assistant-bubble { background:#1e293b; border:1px solid #334155; color:#e5e7eb; }
#         section[data-testid="stSidebar"] { background:#020617; border-right:1px solid #1f2937; }
#         .stChatInput input { background:#020617 !important; color:white !important; border:1px solid #334155 !important; }
#         </style>
#         """
#     else:
#         return """
#         <style>
#         .stApp { background: #f8fafc; color: #0f172a; }
#         .assistant-bubble { background: #e2e8f0; border:1px solid #cbd5e1; color: #0f172a; }
#         section[data-testid="stSidebar"] { background: #f1f5f9; border-right:1px solid #cbd5e1; }
#         .stChatInput input { background: white !important; color: black !important; border:1px solid #cbd5e1 !important; }
#         .chat-header { -webkit-text-fill-color: #0f172a; background: none; }
#         </style>
#         """

# st.markdown(get_theme_css(), unsafe_allow_html=True)

# # Base CSS (shared)
# st.markdown("""
# <style>
# .chat-header { text-align:center; font-size:34px; font-weight:800; margin-bottom:10px; }
# .chat-message { display:flex; align-items:flex-end; margin:14px 0; animation:fadeIn .25s ease-in; }
# @keyframes fadeIn { from {opacity:0; transform:translateY(6px);} to {opacity:1; transform:translateY(0);} }
# .avatar { font-size:22px; margin:0 8px; }
# .user { justify-content:flex-end; }
# .assistant { justify-content:flex-start; }
# .bubble { max-width:70%; padding:14px 18px; border-radius:18px; font-size:15px; line-height:1.6; }
# .user-bubble { background:linear-gradient(135deg,#2563eb,#7c3aed); color:white; border-bottom-right-radius:6px; }
# .assistant-bubble { border-bottom-left-radius:6px; }
# .stButton>button { border-radius:12px; }
# .chat-controls { display:flex; gap:5px; margin-top:5px; }
# .chat-controls button { background:transparent; border:none; color:#aaa; cursor:pointer; font-size:12px; }
# .chat-controls button:hover { color:white; }
# </style>
# """, unsafe_allow_html=True)

# # ===== SIDEBAR =====
# with st.sidebar:
#     st.title("✨ Nova Advanced")

#     # User name (simple memory)
#     st.session_state.user_name = st.text_input("Your name", value=st.session_state.user_name)
#     if st.button("Remember me"):
#         for cid in st.session_state.chats:
#             st.session_state.chats[cid]["memory"]["user_name"] = st.session_state.user_name
#         st.success("Name saved to memory!")

#     st.divider()

#     # New chat
#     if st.button("➕ New Chat", use_container_width=True):
#         new_chat()
#         st.rerun()

#     st.divider()

#     # Chat list with rename/delete
#     for cid, data in sorted(
#         st.session_state.chats.items(),
#         key=lambda x: x[1]["created"],
#         reverse=True
#     ):
#         col1, col2, col3 = st.columns([6, 1, 1])
#         with col1:
#             if st.button(data["title"], key=f"chat_{cid}", use_container_width=True):
#                 st.session_state.current = cid
#                 st.rerun()
#         with col2:
#             if st.button("✏️", key=f"rename_{cid}"):
#                 st.session_state[f"renaming_{cid}"] = True
#         with col3:
#             if st.button("🗑️", key=f"delete_{cid}"):
#                 delete_chat(cid)

#         if st.session_state.get(f"renaming_{cid}", False):
#             new_name = st.text_input("New name", value=data["title"], key=f"rename_input_{cid}")
#             if st.button("Save", key=f"save_rename_{cid}"):
#                 rename_chat(cid, new_name)
#                 st.session_state[f"renaming_{cid}"] = False
#                 st.rerun()

#     st.divider()

#     # Model and temperature
#     model = st.selectbox(
#         "Model",
#         ["llama-3.3-70b-versatile",
#          "mixtral-8x7b-32768",
#          "gemma2-9b-it",
#          "llama-3.1-8b-instant"],
#         index=0
#     )
#     temperature = st.slider("Temperature", 0.0, 2.0, 0.0, 0.1)

#     # Theme toggle
#     theme_choice = st.radio("Theme", ["dark", "light"], index=0 if st.session_state.theme=="dark" else 1)
#     if theme_choice != st.session_state.theme:
#         st.session_state.theme = theme_choice
#         st.rerun()

#     st.divider()

#     # Smart Modes
#     mode = st.selectbox("Mode", ["General", "Study", "Coding", "Interview", "Fun"])
#     if mode == "Study":
#         sys_prompt = "You are a study assistant. Provide clear explanations and examples. Be concise."
#     elif mode == "Coding":
#         sys_prompt = "You are a coding expert. Give clean, efficient code with explanations."
#     elif mode == "Interview":
#         sys_prompt = "You are an interview coach. Ask relevant questions and provide feedback."
#     elif mode == "Fun":
#         sys_prompt = "You are a fun and creative assistant. Be witty and engaging."
#     else:
#         sys_prompt = "You are Nova, a helpful assistant. Be accurate and concise."
#     if st.button("Apply Mode"):
#         update_system_prompt(sys_prompt)
#         st.success(f"Switched to {mode} mode")

#     st.divider()

#     # System prompt editor
#     with st.expander("System Prompt (advanced)"):
#         current_system = st.session_state.chats[st.session_state.current]["system_prompt"]
#         new_system = st.text_area("Edit", value=current_system, height=100)
#         if st.button("Update"):
#             update_system_prompt(new_system)

#     st.divider()

#     # PDF/DOCX upload
#     uploaded_file = st.file_uploader("Upload document (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
#     if uploaded_file is not None:
#         with st.spinner("Extracting text..."):
#             text = extract_text_from_file(uploaded_file)
#             # Store in session or just add to context? We'll add as a user message with note
#             if text:
#                 st.success("Text extracted! You can now ask about it.")
#                 # Option to include in context automatically? We'll put a button to add as context
#                 if st.button("Add document to context"):
#                     add_msg("user", f"[Uploaded document: {uploaded_file.name}]\n\n{text[:2000]}...")  # limit
#                     st.rerun()

#     st.divider()

#     # Export
#     export_format = st.radio("Export format", ["Markdown", "PDF"])
#     if st.button("Export current chat"):
#         if export_format == "Markdown":
#             md = export_chat_as_markdown(st.session_state.current)
#             st.download_button("Download Markdown", data=md, file_name="chat.md", mime="text/markdown")
#         else:
#             pdf_bytes = export_chat_as_pdf(st.session_state.current)
#             st.download_button("Download PDF", data=pdf_bytes, file_name="chat.pdf", mime="application/pdf")

# # ===== MAIN CHAT AREA =====
# chat = st.session_state.chats[st.session_state.current]

# # Display header with memory greeting
# if chat["memory"]["user_name"]:
#     greeting = f"Welcome back, {chat['memory']['user_name']}!"
# else:
#     greeting = "Welcome to Nova!"
# st.markdown(f"<div class='chat-header'>✨ Nova AI <span style='font-size:16px;'>({greeting})</span></div>", unsafe_allow_html=True)

# # ===== DISPLAY MESSAGES =====
# for idx, m in enumerate(chat["messages"]):
#     if m["role"] == "user":
#         st.markdown(f"""
#         <div class="chat-message user">
#             <div class="bubble user-bubble">{m["content"]}</div>
#             <div class="avatar">🧑‍💻</div>
#         </div>
#         """, unsafe_allow_html=True)
#     else:
#         st.markdown(f"""
#         <div class="chat-message assistant">
#             <div class="avatar">✨</div>
#             <div class="bubble assistant-bubble">{m["content"]}</div>
#         </div>
#         """, unsafe_allow_html=True)

# # ===== REGENERATION =====
# if chat["messages"] and chat["messages"][-1]["role"] == "assistant":
#     if st.button("🔄 Regenerate", help="Regenerate last response"):
#         chat["messages"].pop()
#         st.session_state.regenerate_target = st.session_state.current
#         st.rerun()

# # ===== STOP GENERATION =====
# if st.session_state.get("streaming", False):
#     if st.button("⏹️ Stop", key="stop_button"):
#         st.session_state.stop_generation = True

# # ===== INPUT =====
# prompt = st.chat_input("Ask Nova anything...")

# if prompt:
#     add_msg("user", prompt)

#     st.markdown(f"""
#     <div class="chat-message user">
#         <div class="bubble user-bubble">{prompt}</div>
#         <div class="avatar">🧑‍💻</div>
#     </div>
#     """, unsafe_allow_html=True)

#     # Prepare messages for API (no timestamps)
#     messages = get_messages_for_api()

#     placeholder = st.empty()
#     reply = ""
#     st.session_state.streaming = True
#     st.session_state.stop_generation = False

#     try:
#         stream = client.chat.completions.create(
#             model=model,
#             messages=messages,
#             stream=True,
#             temperature=temperature,
#             top_p=0.1,
#             max_tokens=1024
#         )

#         for chunk in stream:
#             if st.session_state.stop_generation:
#                 stream.close()
#                 break
#             if chunk.choices[0].delta.content:
#                 reply += chunk.choices[0].delta.content
#                 placeholder.markdown(f"""
#                 <div class="chat-message assistant">
#                     <div class="avatar">✨</div>
#                     <div class="bubble assistant-bubble">{reply}</div>
#                 </div>
#                 """, unsafe_allow_html=True)

#     finally:
#         st.session_state.streaming = False
#         st.session_state.stop_generation = False

#     if reply.strip():
#         add_msg("assistant", reply)
#     else:
#         st.warning("No response generated. Please try again.")

# # ===== REGENERATION HANDLER =====
# if st.session_state.get("regenerate_target") == st.session_state.current:
#     st.session_state.regenerate_target = None
#     # Find last user message
#     last_user = None
#     for m in reversed(chat["messages"]):
#         if m["role"] == "user":
#             last_user = m["content"]
#             break
#     if last_user:
#         # We have removed the last assistant, so we can re-run generation by simulating input
#         # To avoid duplicating code, we set a flag to regenerate in next run with same prompt
#         st.session_state.regenerate_prompt = last_user
#         st.rerun()

# if "regenerate_prompt" in st.session_state:
#     prompt = st.session_state.pop("regenerate_prompt")
#     # Use same generation code (we'll just call the same block, but careful not to add duplicate user message)
#     # Since we already removed the assistant, the messages list is ready. We'll just call the API again.
#     messages = get_messages_for_api()
#     placeholder = st.empty()
#     reply = ""
#     st.session_state.streaming = True
#     st.session_state.stop_generation = False

#     try:
#         stream = client.chat.completions.create(
#             model=model,
#             messages=messages,
#             stream=True,
#             temperature=temperature,
#             top_p=0.1,
#             max_tokens=1024
#         )

#         for chunk in stream:
#             if st.session_state.stop_generation:
#                 stream.close()
#                 break
#             if chunk.choices[0].delta.content:
#                 reply += chunk.choices[0].delta.content
#                 placeholder.markdown(f"""
#                 <div class="chat-message assistant">
#                     <div class="avatar">✨</div>
#                     <div class="bubble assistant-bubble">{reply}</div>
#                 </div>
#                 """, unsafe_allow_html=True)

#     finally:
#         st.session_state.streaming = False
#         st.session_state.stop_generation = False

#     if reply.strip():
#         add_msg("assistant", reply)

# # Footer
# st.markdown(
#     "<center style='opacity:.4;margin-top:20px'>Nova Advanced • Powered by Groq • Memory • PDF Chat • Themes • Export</center>",
#     unsafe_allow_html=True
# )


















import streamlit as st
from groq import Groq
import time
from datetime import datetime, date
import json
import os
import pdfplumber
from docx import Document
from fpdf import FPDF
import re
import shutil
import requests

# ===== CONFIG =====
st.set_page_config(
    page_title="Nova Advanced",
    page_icon="✨",
    layout="wide"
)

import os
# ===== GROQ API KEY =====
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "your_groq_api_key_here")
client = Groq(api_key=GROQ_API_KEY)

# ===== CONSTANTS =====
CHATS_FILE = "chats.json"
BACKUP_FILE = "backup.json"
BOOKMARKS_FILE = "bookmarks.json"
MEMORY_FILE = "memory.json"
USAGE_FILE = "usage.json"
MAX_CHATS = 50
MAX_HISTORY = 20
MAX_CHUNK_SIZE = 1000
MAX_CHUNKS_TO_USE = 3

# ===== PERSISTENCE =====
def load_json(file, default):
    if os.path.exists(file):
        with open(file, "r", encoding="utf-8") as f:
            return json.load(f)
    return default

def save_json(file, data):
    with open(file, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def load_chats():
    return load_json(CHATS_FILE, {})

def save_chats():
    save_json(CHATS_FILE, st.session_state.chats)
    shutil.copy(CHATS_FILE, BACKUP_FILE)

def load_bookmarks():
    return load_json(BOOKMARKS_FILE, [])

def save_bookmarks():
    save_json(BOOKMARKS_FILE, st.session_state.bookmarks)

def load_memory():
    return load_json(MEMORY_FILE, {"notes": [], "facts": []})

def save_memory():
    save_json(MEMORY_FILE, st.session_state.memory)

def load_usage():
    return load_json(USAGE_FILE, {"daily": {}, "total_messages": 0})

def save_usage():
    save_json(USAGE_FILE, st.session_state.usage)

def delete_oldest_chat():
    if len(st.session_state.chats) > MAX_CHATS:
        oldest_cid = min(
            st.session_state.chats.keys(),
            key=lambda cid: st.session_state.chats[cid]["created"]
        )
        del st.session_state.chats[oldest_cid]

# ===== SESSION STATE INIT =====
# FIX: Use a single 'initialized' flag so this only runs ONCE per session,
# not on every Streamlit rerun (which caused state corruption + duplicates).
def init_session_state():
    if "initialized" in st.session_state:
        return

    st.session_state.chats = load_chats()
    if not st.session_state.chats:
        cid = str(int(time.time()))
        st.session_state.chats = {
            cid: {
                "title": "New Chat",
                "created": datetime.now().isoformat(),
                "system_prompt": (
                    "You are Nova, a highly accurate and factual AI assistant. "
                    "Think step-by-step. Give concise, truthful answers. "
                    "If unsure, say you don't know."
                ),
                "messages": [],
                "memory": {"user_name": "User"},
                "document_text": None,
                "document_chunks": [],
                "title_generated": False
            }
        }
        st.session_state.current = cid
    else:
        for cid, chat in st.session_state.chats.items():
            if "memory" not in chat:
                chat["memory"] = {"user_name": "User"}
            if "document_text" not in chat:
                chat["document_text"] = None
            if "document_chunks" not in chat:
                chat["document_chunks"] = []
            if "title_generated" not in chat:
                chat["title_generated"] = False
        st.session_state.current = next(iter(st.session_state.chats))

    st.session_state.stop_generation = False
    st.session_state.do_regenerate = False
    st.session_state.theme = "dark"
    st.session_state.user_name = (
        st.session_state.chats[st.session_state.current]["memory"].get("user_name", "User")
    )
    st.session_state.document_context_enabled = False
    st.session_state.web_mode = False
    st.session_state.emotion = None
    st.session_state.bookmarks = load_bookmarks()
    st.session_state.memory = load_memory()
    st.session_state.usage = load_usage()

    today_str = date.today().isoformat()
    if today_str not in st.session_state.usage["daily"]:
        st.session_state.usage["daily"][today_str] = {
            "messages": 0,
            "questions": 0,
            "time_seconds": 0
        }
    st.session_state.session_start = time.time()
    st.session_state.initialized = True


init_session_state()


# ===== HELPER FUNCTIONS =====
def new_chat():
    cid = str(int(time.time()))
    st.session_state.chats[cid] = {
        "title": "New Chat",
        "created": datetime.now().isoformat(),
        "system_prompt": (
            "You are Nova, a highly accurate and factual AI assistant. "
            "Think step-by-step. Give concise, truthful answers. "
            "If unsure, say you don't know."
        ),
        "messages": [],
        "memory": {"user_name": st.session_state.user_name},
        "document_text": None,
        "document_chunks": [],
        "title_generated": False
    }
    st.session_state.current = cid
    delete_oldest_chat()
    save_chats()


def delete_chat(cid):
    if cid in st.session_state.chats:
        del st.session_state.chats[cid]
    if st.session_state.current == cid:
        if st.session_state.chats:
            st.session_state.current = next(iter(st.session_state.chats))
        else:
            new_chat()
    save_chats()
    st.rerun()


def rename_chat(cid, new_title):
    if cid in st.session_state.chats:
        st.session_state.chats[cid]["title"] = new_title
        save_chats()


def add_msg(role, content):
    """
    FIX: Add a message only if it's not a duplicate of the last message.
    This is the primary guard against double-saving assistant replies.
    """
    chat = st.session_state.chats[st.session_state.current]
    msgs = chat["messages"]

    # Strict duplicate check: same role AND same content as last msg
    if msgs and msgs[-1]["role"] == role and msgs[-1]["content"] == content:
        return

    msgs.append({
        "role": role,
        "content": content,
        "timestamp": datetime.now().isoformat()
    })

    today_str = date.today().isoformat()
    if today_str not in st.session_state.usage["daily"]:
        st.session_state.usage["daily"][today_str] = {
            "messages": 0, "questions": 0, "time_seconds": 0
        }
    st.session_state.usage["daily"][today_str]["messages"] += 1
    if role == "user":
        st.session_state.usage["daily"][today_str]["questions"] += 1
    st.session_state.usage["total_messages"] += 1
    save_usage()

    if role == "user" and not chat["title_generated"]:
        user_msgs = [m for m in msgs if m["role"] == "user"]
        if len(user_msgs) >= 2:
            generate_chat_title(st.session_state.current)

    save_chats()


def generate_chat_title(cid):
    chat = st.session_state.chats[cid]
    conversation = "\n".join(
        [f"{m['role']}: {m['content']}" for m in chat["messages"][:4]]
    )
    prompt = (
        f"Based on this conversation, generate a very short title (max 3 words) "
        f"that summarizes the topic:\n\n{conversation}\n\nTitle:"
    )
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=10
        )
        title = response.choices[0].message.content.strip().strip('"')
        if title:
            chat["title"] = title[:50]
            chat["title_generated"] = True
            save_chats()
    except Exception as e:
        print(f"Title generation failed: {e}")


def detect_emotion(text):
    prompt = (
        "Classify the emotion of the following user message into one word: "
        "happy, sad, angry, stressed, confused, curious, excited, or neutral. "
        f"Only output the emotion word.\n\nMessage: {text}\n\nEmotion:"
    )
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=10
        )
        emotion = response.choices[0].message.content.strip().lower()
        valid = ["happy", "sad", "angry", "stressed", "confused", "curious", "excited", "neutral"]
        return emotion if emotion in valid else "neutral"
    except:
        return "neutral"


def detect_web_intent(query):
    """
    Detect if a query needs weather or news data, and extract the topic/city.
    Returns: ("weather", city) | ("news", topic) | (None, None)
    """
    q = query.lower()

    # Weather intent keywords
    weather_keywords = ["weather", "temperature", "forecast", "humid", "raining",
                        "sunny", "cloudy", "wind", "hot outside", "cold outside",
                        "what's it like in", "how's the weather", "degrees"]
    is_weather = any(kw in q for kw in weather_keywords)

    # News intent keywords
    news_keywords = ["news", "latest", "headlines", "happening", "today in",
                     "recent", "update", "breaking", "current events"]
    is_news = any(kw in q for kw in news_keywords)

    if is_weather:
        # Extract city: remove noise words to isolate city name
        noise = ["weather", "what", "is", "the", "in", "at", "for", "today",
                 "tomorrow", "like", "how", "whats", "what's", "tell", "me",
                 "current", "now", "temperature", "forecast", "degrees", "please"]
        words = [w for w in re.findall(r"[a-z]+", q) if w not in noise]
        city = " ".join(words).strip() if words else "London"
        return ("weather", city)

    if is_news:
        # Extract topic: strip generic news words to get the subject
        noise = ["news", "latest", "tell", "me", "about", "the", "what", "is",
                 "are", "any", "recent", "updates", "headlines", "today", "show"]
        words = [w for w in re.findall(r"[a-z]+", q) if w not in noise]
        topic = " ".join(words).strip() if words else query
        return ("news", topic)

    return (None, None)


def web_search(query, search_type="news"):
    NEWS_API_KEY = "5167a3decb8b4fb78fcf7ac1446e5e6e"
    WEATHER_API_KEY = "6890236208c02c0adf36ba6fdbab1712"

    try:
        if search_type == "news":
            # Use full query as topic for better results
            topic = query.strip() or "top news"
            url = "https://newsapi.org/v2/everything"
            params = {
                "q": topic,
                "apiKey": NEWS_API_KEY,
                "pageSize": 5,
                "sortBy": "relevancy",
                "language": "en"
            }
            response = requests.get(url, params=params, timeout=10)
            data = response.json()
            if data.get("status") == "ok" and data.get("articles"):
                articles = data["articles"][:3]
                result = f"REAL-TIME NEWS DATA (fetched live right now) about '{topic}':\n\n"
                for i, article in enumerate(articles, 1):
                    result += f"{i}. {article['title']}\n"
                    if article.get("description"):
                        result += f"   Summary: {article['description']}\n"
                    result += f"   Source: {article.get('source', {}).get('name', 'Unknown')}\n"
                    result += f"   URL: {article['url']}\n\n"
                return result
            else:
                return f"NEWS API returned no results for '{topic}'. Status: {data.get('status')} | Code: {data.get('code')}"

        elif search_type == "weather":
            city = query.strip() if query.strip() else "London"
            url = "https://api.openweathermap.org/data/2.5/weather"
            params = {"q": city, "appid": WEATHER_API_KEY, "units": "metric"}
            response = requests.get(url, params=params, timeout=10)
            data = response.json()
            if data.get("cod") == 200:
                return (
                    f"REAL-TIME WEATHER DATA (fetched live right now) for {data['name']}, {data['sys']['country']}:\n"
                    f"Condition: {data['weather'][0]['description'].capitalize()}\n"
                    f"Temperature: {data['main']['temp']}°C (feels like {data['main']['feels_like']}°C)\n"
                    f"Min/Max: {data['main']['temp_min']}°C / {data['main']['temp_max']}°C\n"
                    f"Humidity: {data['main']['humidity']}%\n"
                    f"Wind Speed: {data['wind']['speed']} m/s\n"
                    f"Visibility: {data.get('visibility', 'N/A')} meters"
                )
            else:
                return f"WEATHER API error for city '{city}': {data.get('message', 'Unknown error')}. Try a different city name."

    except requests.exceptions.RequestException as e:
        return f"NETWORK ERROR fetching real-time data: {str(e)}"
    except Exception as e:
        return f"ERROR fetching real-time data: {str(e)}"


def retrieve_memory(query):
    results = []
    for note in st.session_state.memory["notes"]:
        if any(word in note.lower() for word in query.lower().split()):
            results.append(f"Note: {note}")
    for fact in st.session_state.memory["facts"]:
        if any(word in fact.lower() for word in query.lower().split()):
            results.append(f"Fact: {fact}")
    return results


def add_to_memory(content, category="notes"):
    if category == "notes":
        st.session_state.memory["notes"].append(content)
    elif category == "facts":
        st.session_state.memory["facts"].append(content)
    save_memory()


def update_system_prompt(new_prompt):
    st.session_state.chats[st.session_state.current]["system_prompt"] = new_prompt
    save_chats()


# ===== DOCUMENT PROCESSING =====
@st.cache_data
def extract_text_from_file(uploaded_file):
    text = ""
    if uploaded_file.type == "application/pdf":
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif uploaded_file.type == "text/plain":
        text = uploaded_file.read().decode("utf-8")
    return text


def chunk_text(text, chunk_size=MAX_CHUNK_SIZE):
    words = text.split()
    chunks = []
    current_chunk = []
    current_len = 0
    for word in words:
        word_len = len(word) + 1
        if current_len + word_len > chunk_size and current_chunk:
            chunks.append(" ".join(current_chunk))
            overlap = max(1, len(current_chunk) // 4)
            current_chunk = current_chunk[-overlap:]
            current_len = sum(len(w) + 1 for w in current_chunk)
        current_chunk.append(word)
        current_len += word_len
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks


def get_relevant_chunks(query, chunks, top_n=MAX_CHUNKS_TO_USE):
    if not chunks:
        return []
    stopwords = {
        "the", "a", "an", "is", "are", "was", "were", "in", "on",
        "at", "to", "for", "of", "with", "by", "and", "or", "but"
    }
    query_words = set(re.findall(r'\w+', query.lower())) - stopwords
    if not query_words:
        return chunks[:top_n]
    scores = []
    for chunk in chunks:
        chunk_words = set(re.findall(r'\w+', chunk.lower()))
        scores.append(len(query_words.intersection(chunk_words)))
    top_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:top_n]
    return [chunks[i] for i in top_indices if scores[i] > 0]


def get_messages_for_api(user_input=None):
    chat = st.session_state.chats[st.session_state.current]
    messages = []

    # Document context
    if st.session_state.document_context_enabled and chat.get("document_text"):
        last_user = user_input or (
            chat["messages"][-1]["content"]
            if chat["messages"] and chat["messages"][-1]["role"] == "user"
            else ""
        )
        if last_user and chat.get("document_chunks"):
            relevant = get_relevant_chunks(last_user, chat["document_chunks"])
            if relevant:
                doc_context = "Relevant parts of the uploaded document:\n\n" + "\n\n---\n\n".join(relevant)
                messages.append({"role": "system", "content": doc_context})
            elif chat["document_chunks"]:
                messages.append({
                    "role": "system",
                    "content": "Document content (first part):\n\n" + chat["document_chunks"][0]
                })

    # Memory context
    if user_input:
        memory_hits = retrieve_memory(user_input)
        if memory_hits:
            messages.append({
                "role": "system",
                "content": "Relevant information from your personal memory:\n" + "\n".join(memory_hits)
            })

    # Emotion context
    if st.session_state.emotion and st.session_state.emotion != "neutral":
        messages.append({
            "role": "system",
            "content": (
                f"The user seems {st.session_state.emotion}. "
                "Respond with empathy and adapt your tone accordingly."
            )
        })

    # System prompt — append real-time instruction when web mode is on
    base_prompt = chat["system_prompt"]
    if st.session_state.web_mode:
        base_prompt += (
            "\n\nIMPORTANT: When the conversation contains a message starting with "
            "'REAL-TIME WEATHER DATA' or 'REAL-TIME NEWS DATA', you MUST use ONLY "
            "that data to answer. Do NOT say you lack real-time access. "
            "Do NOT suggest checking other websites. Present the provided data "
            "clearly and directly as the answer."
        )
    messages.append({"role": "system", "content": base_prompt})

    # Chat history
    for m in chat["messages"][-MAX_HISTORY:]:
        messages.append({"role": m["role"], "content": m["content"]})

    return messages


# ===== EXPORT FUNCTIONS =====
def export_chat_as_markdown(cid):
    chat = st.session_state.chats[cid]
    lines = [f"# {chat['title']}\n", f"*Created: {chat['created']}*\n"]
    for msg in chat["messages"]:
        role = "**User**" if msg["role"] == "user" else "**Nova**"
        ts = msg.get("timestamp", "")[:16]
        lines.append(f"\n{role} ({ts}): {msg['content']}\n")
    return "\n".join(lines)


def export_chat_as_pdf(cid):
    chat = st.session_state.chats[cid]
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=chat['title'], ln=1, align='C')
    pdf.ln(10)
    for msg in chat["messages"]:
        role = "User: " if msg["role"] == "user" else "Nova: "
        ts = msg.get("timestamp", "")[:16]
        pdf.multi_cell(0, 10, txt=f"{role}({ts}) {msg['content']}")
        pdf.ln(5)
    return pdf.output(dest='S').encode('latin1')


# ===== STREAMING HELPER =====
def run_stream(messages_for_api, model, temperature):
    """
    FIX: Stream response into a placeholder for live preview,
    then CLEAR the placeholder before returning.
    This prevents the streamed text from persisting on screen
    and being rendered again when st.rerun() draws history.
    """
    placeholder = st.empty()
    reply = ""
    st.session_state.stop_generation = False

    try:
        stream = client.chat.completions.create(
            model=model,
            messages=messages_for_api,
            stream=True,
            temperature=temperature,
            top_p=0.1,
            max_tokens=1024
        )
        for chunk in stream:
            if st.session_state.get("stop_generation", False):
                stream.close()
                break
            delta = chunk.choices[0].delta.content
            if delta:
                reply += delta
                # Show live streaming preview with cursor
                placeholder.markdown(f"""
                <div class="chat-message assistant">
                    <div class="avatar">✨</div>
                    <div class="bubble assistant-bubble">{reply}▌</div>
                </div>
                """, unsafe_allow_html=True)
    finally:
        # CRITICAL FIX: Clear the streaming placeholder.
        # After st.rerun(), the saved message in history will render instead.
        # Without this, BOTH the placeholder AND the history message show = duplicate.
        placeholder.empty()

    return reply.strip()


# ===== THEME CSS =====
def get_theme_css():
    if st.session_state.theme == "dark":
        return """
        <style>
        .stApp { background: linear-gradient(135deg, #0f172a, #020617); color: white; }
        .assistant-bubble { background:#1e293b; border:1px solid #334155; color:#e5e7eb; }
        section[data-testid="stSidebar"] { background:#020617; border-right:1px solid #1f2937; }
        .stChatInput input { background:#020617 !important; color:white !important; border:1px solid #334155 !important; }
        </style>
        """
    else:
        return """
        <style>
        .stApp { background: #f8fafc; color: #0f172a; }
        .assistant-bubble { background: #e2e8f0; border:1px solid #cbd5e1; color: #0f172a; }
        section[data-testid="stSidebar"] { background: #f1f5f9; border-right:1px solid #cbd5e1; }
        .stChatInput input { background: white !important; color: black !important; border:1px solid #cbd5e1 !important; }
        .chat-header { -webkit-text-fill-color: #0f172a; background: none; }
        </style>
        """


st.markdown(get_theme_css(), unsafe_allow_html=True)

st.markdown("""
<style>
.chat-header {
    text-align:center; font-size:34px; font-weight:800; margin-bottom:10px;
}
.chat-message {
    display:flex; align-items:flex-end; margin:14px 0;
    animation:fadeIn .25s ease-in;
}
@keyframes fadeIn {
    from {opacity:0; transform:translateY(6px);}
    to {opacity:1; transform:translateY(0);}
}
.avatar { font-size:22px; margin:0 8px; }
.user { justify-content:flex-end; }
.assistant { justify-content:flex-start; }
.bubble {
    max-width:70%; padding:14px 18px; border-radius:18px;
    font-size:15px; line-height:1.6; position: relative;
}
.user-bubble {
    background:linear-gradient(135deg,#2563eb,#7c3aed);
    color:white; border-bottom-right-radius:6px;
}
.assistant-bubble { border-bottom-left-radius:6px; }
.timestamp { font-size: 10px; opacity: 0.6; margin-top: 5px; text-align: right; }
.stButton>button { border-radius:12px; }
</style>
""", unsafe_allow_html=True)


# ===== SIDEBAR =====
with st.sidebar:
    st.title("✨ Nova Advanced")

    st.session_state.user_name = st.text_input(
        "Your name", value=st.session_state.user_name
    )
    if st.button("Remember me"):
        for cid in st.session_state.chats:
            st.session_state.chats[cid]["memory"]["user_name"] = st.session_state.user_name
        save_chats()
        st.success("Name saved to memory!")

    st.divider()

    if st.button("➕ New Chat", use_container_width=True):
        new_chat()
        st.rerun()

    st.divider()

    for cid, data in sorted(
        st.session_state.chats.items(),
        key=lambda x: x[1]["created"],
        reverse=True
    ):
        col1, col2, col3 = st.columns([6, 1, 1])
        with col1:
            if st.button(data["title"], key=f"chat_{cid}", use_container_width=True):
                st.session_state.current = cid
                st.rerun()
        with col2:
            if st.button("✏️", key=f"rename_{cid}"):
                st.session_state[f"renaming_{cid}"] = True
        with col3:
            if st.button("🗑️", key=f"delete_{cid}"):
                delete_chat(cid)

        if st.session_state.get(f"renaming_{cid}", False):
            new_name = st.text_input(
                "New name", value=data["title"], key=f"rename_input_{cid}"
            )
            if st.button("Save", key=f"save_rename_{cid}"):
                rename_chat(cid, new_name)
                st.session_state[f"renaming_{cid}"] = False
                st.rerun()

    st.divider()

    model = st.selectbox(
        "Model",
        [
            "llama-3.3-70b-versatile",
            "mixtral-8x7b-32768",
            "gemma2-9b-it",
            "llama-3.1-8b-instant"
        ],
        index=0
    )
    temperature = st.slider("Temperature", 0.0, 2.0, 0.0, 0.1)

    theme_choice = st.radio(
        "Theme", ["dark", "light"],
        index=0 if st.session_state.theme == "dark" else 1
    )
    if theme_choice != st.session_state.theme:
        st.session_state.theme = theme_choice
        st.rerun()

    st.divider()

    mode = st.selectbox("Mode", ["General", "Study", "Coding", "Interview", "Fun"])
    mode_prompts = {
        "General": "You are Nova, a helpful assistant. Be accurate and concise.",
        "Study":   "You are a study assistant. Provide clear explanations and examples. Be concise.",
        "Coding":  "You are a coding expert. Give clean, efficient code with explanations.",
        "Interview": "You are an interview coach. Ask relevant questions and provide feedback.",
        "Fun":     "You are a fun and creative assistant. Be witty and engaging."
    }
    if st.button("Apply Mode"):
        update_system_prompt(mode_prompts[mode])
        st.success(f"Switched to {mode} mode")

    st.divider()

    with st.expander("System Prompt (advanced)"):
        current_system = st.session_state.chats[st.session_state.current]["system_prompt"]
        new_system = st.text_area("Edit", value=current_system, height=100)
        if st.button("Update Prompt"):
            update_system_prompt(new_system)
            st.success("Prompt updated!")

    st.divider()

    st.subheader("📄 Document Q&A")
    uploaded_file = st.file_uploader(
        "Upload PDF, DOCX, or TXT", type=["pdf", "docx", "txt"]
    )
    if uploaded_file is not None:
        with st.spinner("Extracting and chunking text..."):
            text = extract_text_from_file(uploaded_file)
            if text:
                st.session_state.chats[st.session_state.current]["document_text"] = text
                chunks = chunk_text(text)
                st.session_state.chats[st.session_state.current]["document_chunks"] = chunks
                save_chats()
                st.success(f"Document loaded! ({len(text)} characters, {len(chunks)} chunks)")
            else:
                st.error("Could not extract text.")

    if st.session_state.chats[st.session_state.current].get("document_text"):
        st.session_state.document_context_enabled = st.checkbox(
            "🔍 Use document as context for Q&A",
            value=st.session_state.document_context_enabled
        )
        if st.button("🗑️ Clear document"):
            st.session_state.chats[st.session_state.current]["document_text"] = None
            st.session_state.chats[st.session_state.current]["document_chunks"] = []
            st.session_state.document_context_enabled = False
            save_chats()
            st.rerun()
    else:
        st.info("No document loaded in this chat.")
        st.session_state.document_context_enabled = False

    st.divider()

    with st.expander("📝 Resume & SOP Helper"):
        resume_file = st.file_uploader(
            "Upload Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"], key="resume"
        )
        if resume_file:
            resume_text = extract_text_from_file(resume_file)
            if st.button("Improve Resume"):
                with st.spinner("AI is improving your resume..."):
                    r_prompt = (
                        "Improve the following resume: correct grammar, suggest better phrasing, "
                        "format it nicely. Keep the original information but enhance it.\n\n"
                        + resume_text[:5000]
                    )
                    r_response = client.chat.completions.create(
                        model=model,
                        messages=[{"role": "user", "content": r_prompt}],
                        temperature=0.3,
                        max_tokens=2000
                    )
                    improved = r_response.choices[0].message.content
                    st.text_area("Improved Resume", improved, height=300)
                    st.download_button(
                        "Download as TXT", improved, file_name="improved_resume.txt"
                    )

    with st.expander("📚 Smart Notes Maker"):
        pdf_file = st.file_uploader(
            "Upload PDF for notes", type=["pdf"], key="notes_pdf"
        )
        if pdf_file:
            notes_text = extract_text_from_file(pdf_file)
            if st.button("Generate Notes"):
                with st.spinner("Creating revision notes..."):
                    n_prompt = (
                        "Summarize the following text and extract key points in a bullet list "
                        "suitable for revision notes:\n\n" + notes_text[:5000]
                    )
                    n_response = client.chat.completions.create(
                        model=model,
                        messages=[{"role": "user", "content": n_prompt}],
                        temperature=0.3,
                        max_tokens=1500
                    )
                    notes = n_response.choices[0].message.content
                    st.markdown(notes)
                    st.download_button(
                        "Download Notes", notes, file_name="revision_notes.txt"
                    )

    st.divider()

    with st.expander("⭐ Bookmarked Messages"):
        if st.session_state.bookmarks:
            for i, bm in enumerate(st.session_state.bookmarks):
                st.markdown(f"**{bm['role']}** ({bm['timestamp'][:16]}):")
                st.caption(bm['content'][:100] + "...")
                if st.button("Delete", key=f"del_bm_{i}"):
                    st.session_state.bookmarks.pop(i)
                    save_bookmarks()
                    st.rerun()
                st.divider()
        else:
            st.info("No bookmarks yet. Click ⭐ on any message to save.")

    st.divider()

    with st.expander("🧠 Offline Memory"):
        st.subheader("Your Notes")
        new_note = st.text_area("Add a note")
        if st.button("Save Note"):
            if new_note.strip():
                add_to_memory(new_note, "notes")
                st.success("Note saved!")

        st.subheader("Facts / Study Materials")
        new_fact = st.text_area("Add a fact")
        if st.button("Save Fact"):
            if new_fact.strip():
                add_to_memory(new_fact, "facts")
                st.success("Fact saved!")

        if st.button("View All Memory"):
            st.write(st.session_state.memory)

    st.divider()

    st.session_state.web_mode = st.checkbox(
        "🌐 Real-Time Web Mode", value=st.session_state.web_mode
    )
    if st.session_state.web_mode:
        st.info("Web mode active – news & weather searches enabled.")

    st.divider()

    with st.expander("📊 Daily Usage"):
        today_str = date.today().isoformat()
        usage_today = st.session_state.usage["daily"].get(
            today_str, {"messages": 0, "questions": 0, "time_seconds": 0}
        )
        st.write(f"Messages today: {usage_today['messages']}")
        st.write(f"Questions asked: {usage_today['questions']}")
        elapsed = int(time.time() - st.session_state.session_start)
        st.write(f"Time this session: {elapsed // 60} min {elapsed % 60} sec")
        total_time = usage_today["time_seconds"] + elapsed
        st.write(f"Total time today: {total_time // 60} min {total_time % 60} sec")

    st.divider()

    export_format = st.radio("Export format", ["Markdown", "PDF"])
    if st.button("Export current chat"):
        if export_format == "Markdown":
            md = export_chat_as_markdown(st.session_state.current)
            st.download_button(
                "Download Markdown", data=md, file_name="chat.md", mime="text/markdown"
            )
        else:
            pdf_bytes = export_chat_as_pdf(st.session_state.current)
            st.download_button(
                "Download PDF", data=pdf_bytes, file_name="chat.pdf", mime="application/pdf"
            )

    st.divider()
    st.caption(f"Total messages: {st.session_state.usage['total_messages']}")


# ===== MAIN CHAT AREA =====
chat = st.session_state.chats[st.session_state.current]

greeting = (
    f"Welcome back, {chat['memory']['user_name']}!"
    if chat["memory"].get("user_name") else "Welcome to Nova!"
)
if st.session_state.emotion and st.session_state.emotion != "neutral":
    greeting += f" (Feeling {st.session_state.emotion})"

st.markdown(
    f"<div class='chat-header'>✨ Nova AI "
    f"<span style='font-size:16px;'>({greeting})</span></div>",
    unsafe_allow_html=True
)

# ===== RENDER SAVED MESSAGES =====
# FIX: Only render messages from session state history here.
# The streaming placeholder (in run_stream) is cleared before rerun,
# so there is never a case where both the placeholder and history show the same message.
for idx, m in enumerate(chat["messages"]):
    ts = m.get("timestamp", "")[:16]
    if m["role"] == "user":
        st.markdown(f"""
        <div class="chat-message user">
            <div class="bubble user-bubble">
                {m["content"]}
                <div class="timestamp">{ts}</div>
            </div>
            <div class="avatar">🧑‍💻</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="chat-message assistant">
            <div class="avatar">✨</div>
            <div class="bubble assistant-bubble">
                {m["content"]}
                <div class="timestamp">{ts}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        col1, col2, col3 = st.columns([8, 1, 1])
        with col2:
            if st.button("⭐", key=f"bookmark_{idx}", help="Bookmark this message"):
                st.session_state.bookmarks.append({
                    "role": "Nova",
                    "content": m["content"],
                    "timestamp": m.get("timestamp", datetime.now().isoformat()),
                    "chat_title": chat["title"]
                })
                save_bookmarks()
                st.success("Bookmarked!")
        with col3:
            if st.button("📋", key=f"copy_{idx}", help="Show copyable text"):
                st.code(m["content"], language="text")

# Regenerate button — only show after an assistant reply
if chat["messages"] and chat["messages"][-1]["role"] == "assistant":
    if st.button("🔄 Regenerate last response"):
        # Remove last assistant message, set flag, rerun
        chat["messages"].pop()
        save_chats()
        st.session_state.do_regenerate = True
        st.rerun()

# Stop button placeholder (shown during streaming via session state)
stop_placeholder = st.empty()

# ===== HANDLE NEW USER MESSAGE =====
prompt = st.chat_input("Ask Nova anything...")

if prompt:
    # Step 1: Detect emotion
    st.session_state.emotion = detect_emotion(prompt)

    # Step 2: Fetch real-time web data if web mode is on
    # FIX: Inject data as a USER-role message right before the question,
    # NOT as a system message. This forces the LLM to actually read and use it.
    web_context = ""
    if st.session_state.web_mode:
        intent, topic = detect_web_intent(prompt)
        if intent == "weather":
            web_context = web_search(topic, "weather")
        elif intent == "news":
            web_context = web_search(topic, "news")

    # Step 3: Save user message to history
    add_msg("user", prompt)

    # Step 4: Build API messages
    api_messages = get_messages_for_api(user_input=prompt)

    # Step 5: If web data exists, inject it as a SYSTEM message immediately before
    # the last user message (at the end of the list), with a very strong instruction.
    if web_context:
        # Find and wrap the last user message with the real-time data
        # by inserting a high-priority system message just before it
        api_messages.insert(
            len(api_messages) - 1,  # Just before the last (user) message
            {
                "role": "system",
                "content": (
                    f"⚠️ REAL-TIME DATA AVAILABLE — YOU MUST USE THIS:\n\n"
                    f"{web_context}\n\n"
                    f"INSTRUCTIONS: Answer the user's question using ONLY the above "
                    f"real-time data. Do NOT say you don't have access to current info. "
                    f"Do NOT suggest other websites. Just present the data above clearly."
                )
            }
        )

    # Step 6: Show stop button & stream response
    with stop_placeholder:
        if st.button("⏹️ Stop generation", key="stop_btn"):
            st.session_state.stop_generation = True

    reply = run_stream(api_messages, model, temperature)
    stop_placeholder.empty()

    # Step 7: Save assistant reply ONCE to history
    if reply:
        add_msg("assistant", reply)
    else:
        st.warning("No response generated. Please try again.")

    # Step 8: Update usage time
    today_str = date.today().isoformat()
    st.session_state.usage["daily"][today_str]["time_seconds"] = int(
        time.time() - st.session_state.session_start
    )
    save_usage()

    # Step 9: Rerun — placeholder was cleared, so only history renders = no duplicate
    st.rerun()


# ===== HANDLE REGENERATE =====
# FIX: Simple do_regenerate flag instead of complex multi-flag state machine.
if st.session_state.get("do_regenerate"):
    st.session_state.do_regenerate = False

    # Find the last user message
    last_user_msg = None
    for m in reversed(chat["messages"]):
        if m["role"] == "user":
            last_user_msg = m["content"]
            break

    if last_user_msg:
        api_messages = get_messages_for_api(user_input=last_user_msg)

        with stop_placeholder:
            if st.button("⏹️ Stop generation", key="stop_regen_btn"):
                st.session_state.stop_generation = True

        reply = run_stream(api_messages, model, temperature)
        stop_placeholder.empty()

        if reply:
            add_msg("assistant", reply)
        else:
            st.warning("No response generated. Please try again.")

        st.rerun()


# ===== FOOTER =====
st.markdown(
    "<center style='opacity:.4;margin-top:20px'>"
    "Nova Advanced • Persistent Storage • Smart PDF Q&A • "
    "Themes • Export • Bookmarks • Memory • Web Mode • Emotion AI"
    "</center>",
    unsafe_allow_html=True
)


