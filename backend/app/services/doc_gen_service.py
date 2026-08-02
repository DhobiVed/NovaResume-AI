import os
import uuid
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from app.core.config import GENERATED_DIR

class DocGenService:
    def generate_pdf(self, title: str, content: str) -> str:
        filename = f"gen_{uuid.uuid4().hex[:8]}.pdf"
        file_path = os.path.join(GENERATED_DIR, filename)

        # Single page fitting margins (36pt / 0.5 inch)
        doc = SimpleDocTemplate(
            file_path,
            pagesize=letter,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36
        )
        styles = getSampleStyleSheet()

        primary_color = colors.HexColor('#1e40af') # Professional blue accent
        dark_text = colors.HexColor('#0f172a')

        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontSize=18,
            leading=22,
            textColor=primary_color,
            fontName='Helvetica-Bold',
            spaceAfter=6
        )

        section_heading = ParagraphStyle(
            'SectionHead',
            parent=styles['Heading2'],
            fontSize=12,
            leading=15,
            textColor=primary_color,
            fontName='Helvetica-Bold',
            spaceBefore=8,
            spaceAfter=4
        )

        body_style = ParagraphStyle(
            'DocBody',
            parent=styles['Normal'],
            fontSize=9.5,
            leading=13,
            textColor=dark_text,
            fontName='Helvetica',
            spaceAfter=4
        )

        story = [Paragraph(title, title_style), HRFlowable(width="100%", thickness=1.5, color=primary_color, spaceAfter=8)]

        for line in content.split("\n"):
            line_str = line.strip()
            if not line_str:
                story.append(Spacer(1, 4))
                continue

            clean_line = line_str.replace("<", "&lt;").replace(">", "&gt;")

            if line_str.isupper() and len(line_str) < 30:
                # Section heading
                story.append(Paragraph(clean_line, section_heading))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#cbd5e1'), spaceAfter=4))
            elif line_str.startswith("- ") or line_str.startswith("• "):
                bullet_style = ParagraphStyle(
                    'BulletStyle',
                    parent=body_style,
                    leftIndent=12,
                    spaceAfter=2
                )
                story.append(Paragraph(clean_line, bullet_style))
            else:
                story.append(Paragraph(clean_line, body_style))

        doc.build(story)
        return file_path

    def generate_docx(self, title: str, content: str) -> str:
        filename = f"gen_{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(GENERATED_DIR, filename)

        doc = docx.Document()
        doc.add_heading(title, level=0)

        for paragraph_text in content.split("\n"):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())

        doc.save(file_path)
        return file_path

    def generate_txt(self, title: str, content: str) -> str:
        filename = f"gen_{uuid.uuid4().hex[:8]}.txt"
        file_path = os.path.join(GENERATED_DIR, filename)

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"=== {title} ===\n\n{content}")

        return file_path

    def generate_markdown(self, title: str, content: str) -> str:
        filename = f"gen_{uuid.uuid4().hex[:8]}.md"
        file_path = os.path.join(GENERATED_DIR, filename)

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"# {title}\n\n{content}")

        return file_path

doc_gen_service = DocGenService()
